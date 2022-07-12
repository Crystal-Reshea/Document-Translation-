from nltk.tokenize import RegexpTokenizer
from nltk.tokenize import sent_tokenize
from transformers import MarianMTModel, MarianTokenizer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from nltk.tokenize import sent_tokenize
from nltk.tokenize import LineTokenizer
import nltk
from docx import Document
import docx
import re
import math
import torch
import numpy as np
import fugashi

class Translator_Inference:
    def __init__(self, file, source, target): 
            self.file = file
            self.file_type = self.set_type(self.file)
            self.model_path_dict = {
                'english-spanish':"Helsinki-NLP/opus-mt-en-es",
                'spanish-english': "Helsinki-NLP/opus-mt-es-en",
                'french-english': "Helsinki-NLP/opus-mt-fr-en",
                'english-french':"Helsinki-NLP/opus-mt-en-fr",
                'english-japanese':"Helsinki-NLP/opus-tatoeba-en-ja",
                'japanese-english':"Helsinki-NLP/opus-mt-ja-en"
            }
            self.audio_reader = 'patrickvonplaten/wav2vec2-base-100h-with-lm'
            self.file = file
            self.source = source
            self.target = target
            self.keys = []
    
    def set_type(self, file):
        if file.endswith('.docx'): 
            return 'docx'
        elif file.endswith(('.m4a', '.wav','.flac', '.mp3', '.wma', '.aac')):
            return 'audio'
    
    def get_type(self):
        return self.file_type
    
    def get_model(self):
        search = self.source + '-' + self.target
        if search in self.model_path_dict: 
            self.mod_path = self.model_path_dict[search]
            return self.mod_path
        else:
            print("Source and Target languages not yet available.")
    
    def get_text(self, file):
        if self.file_type == 'audio':
            return None
        elif self.file_type == 'docx':
            self.docx = docx.Document(self.file)
            return self.getDocxText(self.docx)
        else:
            print("please insert audio or docx files")
    
    def getDocxText(self, doc):
        fullText = []
        # Reading in paragraphs
        for para in doc.paragraphs:
            fullText.append(para.text)
        # Reading in text from tables
        for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            fullText.append(paragraph.text)
        return '\n'.join(fullText)
    
    def translate_paragraphs(self, paragraphs):
        if torch.cuda.is_available():  
          dev = "cuda"
        else:  
          dev = "cpu" 
        device = torch.device(dev)
        tokenizer = AutoTokenizer.from_pretrained(self.mod_path)
        model = AutoModelForSeq2SeqLM.from_pretrained(self.mod_path)
        model.to(device)
        translated_paragraphs = []
        for paragraph in paragraphs:
            sents = self.get_sentences(paragraph)
            num_sents = len(sents)
            total_words = self.get_total_words(sents)
            if total_words < 209 and self.source != 'japanese':
                # self.keys.extend(paragraph)
                # WAS
                self.keys.extend(sents)
                result = self.translate_chunk(sents)
                translated_paragraphs.extend(result)
            else:
                batch_size = self.get_batch_size(sents, num_sents)
                batches = math.ceil(num_sents / batch_size)
                translated = []
                for i in range(batches):
                    sent_batch = sents[i*batch_size: (i+1)*batch_size]
                    # to_add = [' '.join(sent_batch)]
                    # self.keys += to_add
                    # OR
                    self.keys.extend(sent_batch)
                    # What Happens??
                    model_inputs = tokenizer(sent_batch, return_tensors="pt", padding=True, truncation=True, max_length=500).to(device)
                    with torch.no_grad():
                        translated_batch = model.generate(**model_inputs)
                # og code
                    translated += translated_batch
                translated = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
                translated_paragraphs += [" ".join(translated)]
                # return translated_paragraphs
                #
                # New tryy
        return translated_paragraphs, self.keys
    
    def translate_chunk(self, chunk): 
        # Set up model
        if torch.cuda.is_available():  
          dev = "cuda"
        else:  
          dev = "cpu" 
        device = torch.device(dev)
        tokenizer = AutoTokenizer.from_pretrained(self.mod_path)
        model = AutoModelForSeq2SeqLM.from_pretrained(self.mod_path)
        model.to(device)
        
        # Translate Chunks
        model_inputs = tokenizer(chunk, return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
        with torch.no_grad():
            model_inputs = model.generate(**model_inputs)
        decoded_text = [tokenizer.decode(t, skip_special_tokens=True) for t in model_inputs]
        return decoded_text

    
    def get_sentences(self, paragraph):
        if self.source == 'japanese':
            jpn_sent_tokenize = nltk.RegexpTokenizer(u'[^！？。]*[！？。]')
            return(jpn_sent_tokenize.tokenize(paragraph))
        else:
            sent = sent_tokenize(paragraph)
            return sent
    
    def get_total_words(self, sentences): 
        total = 0
        if self.source == 'japanese':
            for sent in sentences:
                tagger = fugashi.Tagger()
                words = [word.surface for word in tagger(sent)]
                total+= len(words)
        else:
            for sent in sentences: 
                total += len(sent.split(" "))
        return total
    
    def get_batch_size(self, sentences, batch_size):
        if batch_size <= 1 or self.source == 'japanese': 
            return 1
        else:
            batches = math.ceil(len(sentences) / batch_size) 
            for i in range(batches):
                if self.get_total_words(sentences[i*batch_size: (i+1)*batch_size]) > 209: 
                    return self.get_batch_size(sentences, batch_size - 1)
            return batch_size
            
            
    def replace_text_in_paragraph(self, paragraph, key, value):
        change_tracker = False
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)
                    change_tracker = True
            if change_tracker == False:
                paragraph.text = paragraph.text.replace(key, value)
   
    def docx_replace(self, doc_obj, pairs):
        for key, value in pairs.items():
            for paragraph in doc_obj.paragraphs:
                self.replace_text_in_paragraph(paragraph, key, value)
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, key, value)
        return doc_obj
    
    def run_translate(self):
        if self.file_type == 'docx':
            text = self.get_text(self.file)
            lt = LineTokenizer()
            self.get_model()
            paragraphs = lt.tokenize(text)
            # x = self.translate_paragraphs(paragraphs)
            # print(x)
            translated_paragraphs, keys = self.translate_paragraphs(paragraphs)
            values = []
            temp  = self.source
            self.source = self.target
            for tp in translated_paragraphs:
                values.extend(self.get_sentences(tp))
            self.source = temp
            # values = translated_paragraphs
            # print(translated_paragraphs)
            # print(values)
            # print(keys)
            # print(len(values))
            # print(len(keys))
            self.pairs = {keys[i]: values[i] for i in range(len(keys))}
            return self.docx_replace(self.docx, self.pairs)
        elif self.file_type == 'audio':
            return None
        else:
            return("Error: File Type not recognized.")
        
if __name__=='__main__':
    source = 'spanish'
    target = 'english'
    asr = Translator_Inference("NYT_spanish.docx", source, target)
    doc = asr.run_translate()
    print('done')
    doc.save("term_span_engl.docx")