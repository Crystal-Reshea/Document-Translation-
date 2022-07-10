from nltk.tokenize import RegexpTokenizer
from nltk.tokenize import sent_tokenize
from transformers import MarianMTModel, MarianTokenizer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from nltk.tokenize import sent_tokenize
from nltk.tokenize import LineTokenizer
from docx import Document
import docx
import re
import math
import torch
import numpy as np

class Translator_Inference:
    def __init__(self, file, source, target): 
            self.file_type = self.get_type(file)
            self.model_path_dict = {
                'english-spanish':"Helsinki-NLP/opus-mt-en-es",
                'spanish-english': "Helsinki-NLP/opus-mt-es-en",
                'french-english': "Helsinki-NLP/opus-mt-fr-en",
                'english-french':"Helsinki-NLP/opus-mt-en-fr",
                'english-japanese':"Helsinki-NLP/opus-tatoeba-en-ja",
                'japanese-english':"Helsinki-NLP/opus-mt-ja-en"
            }
            self.audio_reader = 'patrickvonplaten/wav2vec2-base-100h-with-lm'
            self.file = file
            self.source = source
            self.target = target
    def get_type(self, file):
        if file.endswith('.docx'): 
            return 'docx'
        elif file.endswith(('.m4a', '.wav','.flac', '.mp3', '.wma', '.aac')):
            return 'audio'
    
    def get_model(self):
        search = self.source + '-' + self.target
        if search in self.model_path_dict: 
            self.mod_path = self.model_path_dict[search]
            return self.mod_path
        else:
            print("Source and Target languages not yet available.")
    
    def get_text(self, file):
        if self.file_type == 'audio':
            return None
        elif self.file_type == 'docx':
            self.docx = docx.Document(self.file)
            return self.getDocxText(self.docx)
        else:
            print("please insert audio or docx files")
    
    def getDocxText(self, doc):
        fullText = []
        # Reading in paragraphs
        for para in doc.paragraphs:
            fullText.append(para.text)
        # Reading in text from tables
        for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            fullText.append(paragraph.text)
        return '\n'.join(fullText)
    
    def translate_paragraphs(self, paragraphs):
        lt = LineTokenizer()

        if torch.cuda.is_available():  
          dev = "cuda"
        else:  
          dev = "cpu" 
        device = torch.device(dev)

        tokenizer = AutoTokenizer.from_pretrained(self.mod_path)

        model = AutoModelForSeq2SeqLM.from_pretrained(self.mod_path)
        model.to(device)

        # Translating sentences
        # Could choose batch size dynamically based on the length of the sentences.
        # Ideally no sentence/batch will be greater than 209 words
        keys = []
        translated_paragraphs = []
        for paragraph in paragraphs:
            batch_size = 4
            sentences = self.get_sentences(paragraph)
            total_words = self.get_total_words(sentences)
            words_per_sent = math.ceil(total_words / len(sentences))
            if total_words > 260 and len(sentences) > 4 : 
                batch_size = self.get_batch_size(total_words, words_per_sent, sentences, batch_size)
            batches = math.ceil(len(sentences) / batch_size)     
            translated = []
            for i in range(batches):
                # selecting the sentences to batch
                sent_batch = sentences[i*batch_size:(i+1)*batch_size]
                keys.extend(sent_batch)
                model_inputs = tokenizer(sent_batch, return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
                with torch.no_grad():
                    translated_batch = model.generate(**model_inputs)
                translated += translated_batch
            translated = [tokenizer.decode(t, skip_special_tokens=True) for t in translated]
            translated_paragraphs += [" ".join(translated)]

        translated_text = "\n".join(translated_paragraphs)
        return translated_paragraphs, keys
    
    def get_sentences(self, paragraph):
        if self.source == 'japanese':
            jp_sent_tokenize = nltk.RegexpTokenizer(u'[^ 「」!?。．）]*[!?。]')
            return(jpn_sent_tokenize.tokenize(paragraph))
        else:
            sent = sent_tokenize(paragraph)
            return sent
    
    def get_total_words(self, sentences):
        total = 0
        for sent in sentences: 
            total += len(sent.split())
        return total
    
    def get_batch_size(self, sentences, batch_size):
        if batch_size <= 1: 
            return 1
        batches = math.ceil(len(sentences) / batch_size) 
        for i in range(batches):
            if self.get_total_words(sentences[i*batch_size: (i+1)*batch_size]) > 260: 
                return get_batch_size(sentence, batch_size - 1)
        return batch_size
            
            
    def replace_text_in_paragraph(self, paragraph, key, value):
        change_tracker = False
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)
                    change_tracker = True
            if change_tracker == False:
                paragraph.text = paragraph.text.replace(key, value)
   
    def docx_replace(self, doc_obj, pairs):
        for key, value in pairs.items():
            for paragraph in doc_obj.paragraphs:
                self.replace_text_in_paragraph(paragraph, key, value)
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, key, value)
        return doc_obj
    
    def run_translate(self):
        if self.file_type == 'docx':
            text = self.get_text(self.file)
            lt = LineTokenizer()
            self.get_model()
            paragraphs = lt.tokenize(text)   
            translated_paragraphs, keys = self.translate_paragraphs(paragraphs)
            values = []
            for p in translated_paragraphs:
                values.extend(sent_tokenize(p))
            self.pairs = {keys[i]: values[i] for i in range(len(keys))}
            return self.docx_replace(self.docx, self.pairs)
        elif self.file_type == 'audio':
            return None
        else:
            return("Error: File Type not recognized.")
    
    
if __name__=='__main__':
    source = 'english'
    target = 'french'
    asr = translator_inference("testing_standard.docx", source, target)
    translation = asr.run_translate()
    print("Translation: " + str(translation))
                